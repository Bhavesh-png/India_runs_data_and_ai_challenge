import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    # Set cell background color
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # Set cell margins (padding) in twentieths of a point (dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_slide_header(doc, title_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run(title_text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229) # Indigo #4F46E5
    
    # Add a thin colored divider line using a bottom border or paragraph spacing
    p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="8" w:color="E5E7EB"/></w:pBdr>')
    p._p.get_or_add_pPr().append(p_border)

def add_bullet(doc, bold_prefix, text_content, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    
    run_bold = p.add_run(bold_prefix + " ")
    run_bold.font.name = 'Calibri'
    run_bold.font.size = Pt(14)
    run_bold.font.bold = True
    run_bold.font.color.rgb = RGBColor(17, 24, 39) # Charcoal
    
    run_text = p.add_run(text_content)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(14)
    run_text.font.color.rgb = RGBColor(55, 65, 81) # Slate

def main():
    doc = docx.Document()
    
    # Set page layout to Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # Set margins (0.5 inches all around for presentation slide feel)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # --- SLIDE 1: Title Slide ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(80)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("SmartHire AI: Candidate Discovery & Ranking")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(36)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(79, 70, 229) # Indigo
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(40)
    run_sub = p_sub.add_run("AI-Powered Talent Matching for Senior AI Engineer (Founding Team)")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(107, 114, 128) # Gray
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("Redrob Data & AI Challenge Solution  |  Technical Overview & Architecture")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(12)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(156, 163, 175)
    
    # --- SLIDE 2: Executive Summary ---
    doc.add_page_break()
    add_slide_header(doc, "1. Executive Summary & The Recruitment Challenge")
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(10)
    p_intro.paragraph_format.space_after = Pt(15)
    run_intro = p_intro.add_run(
        "Standard recruitment tools fail because they focus on keyword matching rather than role understanding. "
        "SmartHire AI acts like an expert recruiter, utilizing semantic similarity and multi-dimensional behavioral signals "
        "to surface candidates who genuinely fit the Founding Team role at Redrob."
    )
    run_intro.font.name = 'Calibri'
    run_intro.font.size = Pt(15)
    run_intro.font.color.rgb = RGBColor(17, 24, 39)
    
    add_bullet(doc, "• Beyond Keywords:", "Matches structural competencies (e.g., building recommendation systems) rather than simple keyword matches (e.g., 'Pinecone' or 'RAG' tags).")
    add_bullet(doc, "• The Founding Team Bar:", "Requires deep ML depth (dense retrieval, evaluation metrics) combined with a scrappy, product-focused shipping attitude.")
    add_bullet(doc, "• Intent & Availability:", "Incorporates actual platform engagement, notice periods, and location alignments into the ranking criteria.")
    add_bullet(doc, "• Honeypot Safeguards:", "Automatically detects and disqualifies simulated/invalid profiles to ensure high shortlist reliability.")
    
    # --- SLIDE 3: System Architecture ---
    doc.add_page_break()
    add_slide_header(doc, "2. System Architecture & Processing Pipeline")
    
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.space_before = Pt(5)
    p_arch.paragraph_format.space_after = Pt(15)
    run_arch = p_arch.add_run("A hybrid, precomputation-based retrieval and scoring architecture designed for maximum performance.")
    run_arch.font.name = 'Calibri'
    run_arch.font.size = Pt(14)
    run_arch.font.italic = True
    
    # Create a simple table to display side-by-side architectural stages
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    
    widths = [Inches(3.2), Inches(3.2), Inches(3.2)]
    stages = [
        ("Offline Precomputation", "Pre-computes candidate embeddings from full profiles (Current Title, Headline, Experience, Summary, Skills, History, Education) using SentenceTransformers all-MiniLM-L6-v2. Restricts token footprint and saves to Disk.", "F4F5F7"),
        ("Semantic Matching", "Loads job description from DOCX and encodes it on demand. Computes matrix dot product of normalized embeddings for lightning-fast (under 5s) candidate-JD cosine similarity.", "EEF2FF"),
        ("Hybrid Multi-Factor Scorer", "Calculates composite scores by combining semantic similarity (60%) and explicit skill alignment (40%), and then applies multipliers based on experience, location, and behavior.", "ECFDF5")
    ]
    
    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = widths[idx]
        title, desc, bg = stages[idx]
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(8)
        crun_title = cp.add_run(title)
        crun_title.font.name = 'Segoe UI'
        crun_title.font.size = Pt(15)
        crun_title.font.bold = True
        crun_title.font.color.rgb = RGBColor(17, 24, 39)
        
        crun_desc = cp.add_run("\n\n" + desc)
        crun_desc.font.name = 'Calibri'
        crun_desc.font.size = Pt(12)
        crun_desc.font.color.rgb = RGBColor(55, 65, 81)

    # --- SLIDE 4: Honeypot & Quality Filters ---
    doc.add_page_break()
    add_slide_header(doc, "3. Profile Quality & Honeypot Detection")
    
    p_hp = doc.add_paragraph()
    p_hp.paragraph_format.space_before = Pt(5)
    p_hp.paragraph_format.space_after = Pt(15)
    run_hp = p_hp.add_run("To protect recruiters from spam, the scoring engine scans candidates for logical contradictions and disqualifies them immediately.")
    run_hp.font.name = 'Calibri'
    run_hp.font.size = Pt(14)
    run_hp.font.italic = True
    
    add_bullet(doc, "• Startup Foundation Conflict:", "Detects candidates claiming to work at startups prior to their official founding date (e.g. Swiggy in 2010 when founded in 2014). Verified against a curated truth table of 30+ startups.")
    add_bullet(doc, "• Experience-Timeline Discrepancy:", "Checks if stated years of experience exceed the actual timeline span of their career history (earliest start date to 2026) by more than 1.5 years.")
    add_bullet(doc, "• Skill Duration Anomalies:", "Flags profiles containing multiple 'expert' or 'advanced' skills with 0 months of duration (indicates randomized or spam keyword stuffing).")
    add_bullet(doc, "• Strict Exclusions:", "Automatically sets the score of any candidate triggering these checks to 0.0, ensuring 0% honeypot rate in the top 100 shortlist.")

    # --- SLIDE 5: Scoring Framework ---
    doc.add_page_break()
    add_slide_header(doc, "4. Multi-Factor Scoring Formulation")
    
    p_score = doc.add_paragraph()
    p_score.paragraph_format.space_before = Pt(5)
    p_score.paragraph_format.space_after = Pt(15)
    run_score = p_score.add_run("Our ranking system applies a transparent, explainable formula to grade candidates:")
    run_score.font.name = 'Calibri'
    run_score.font.size = Pt(14)
    run_score.font.bold = True
    
    # Formula Box
    table_f = doc.add_table(rows=1, cols=1)
    cell_f = table_f.rows[0].cells[0]
    cell_f.width = Inches(9.6)
    set_cell_background(cell_f, "F3F4F6")
    set_cell_margins(cell_f, top=160, bottom=160, left=200, right=200)
    pf = cell_f.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("Final Score = ( [0.6 * Semantic Similarity + 0.4 * Skill Match] * Experience Multiplier * Behavioral Multiplier ) + Activity Bonus")
    rf.font.name = 'Courier New'
    rf.font.size = Pt(13)
    rf.font.bold = True
    rf.font.color.rgb = RGBColor(17, 24, 39)
    
    add_bullet(doc, "• Semantic Similarity (60% weight):", "Cosine similarity between candidate profile embedding and the JD text.", 0)
    add_bullet(doc, "• Skill Match Score (40% weight):", "Evaluates required skills (Embeddings/RAG, Vector DBs, Python, Eval Frameworks) and preferred skills (Fine-tuning, LTR) using proficiency and duration metrics.", 0)
    add_bullet(doc, "• Experience Multiplier:", "Aligns candidate experience with the target 5-9 year curve, applying a penalty for under-experience and a slow decay for over-qualification.", 0)
    add_bullet(doc, "• Behavioral Multiplier:", "Multiplicative discount based on active signals (recruiter response, notice period, location willingness).", 0)

    # --- SLIDE 6: Behavioral Signals & Multipliers ---
    doc.add_page_break()
    add_slide_header(doc, "5. Behavioral Signals & Availability Modifiers")
    
    p_beh = doc.add_paragraph()
    p_beh.paragraph_format.space_before = Pt(5)
    p_beh.paragraph_format.space_after = Pt(15)
    run_beh = p_beh.add_run("A candidate who is technically perfect but unavailable or unresponsive is a poor hire. We calibrate ranks using live indicators:")
    run_beh.font.name = 'Calibri'
    run_beh.font.size = Pt(14)
    run_beh.font.italic = True
    
    add_bullet(doc, "• Recruiter Response Rate:", "Provides a 1.05x boost for highly responsive candidates (>=75%), while heavily penalizing unresponsive ones (<10% response rate receives a 0.55x penalty).")
    add_bullet(doc, "• Notice Period Multipliers:", "Sub-30-day notice period receives a 1.05x speed-joiner boost. Long notice periods (90+ days) receive a 0.65x penalty to favor fast hires.")
    add_bullet(doc, "• Location Alignment & Relocation:", "Local candidates (Pune/Noida/Delhi NCR) get a 1.10x local boost. Tier-1 Indian city candidates willing to relocate get a 1.0x neutral multiplier, while unwilling candidates are penalized (0.70x).")
    add_bullet(doc, "• Platform Activity Bonuses:", "Small additive bonuses (+0.02 to +0.03) are awarded for candidates with high recruiter saves, profile views, or strong GitHub contribution activity.")

    # --- SLIDE 7: Technical Performance & Scalability ---
    doc.add_page_break()
    add_slide_header(doc, "6. Offline Optimization & Scalability")
    
    p_tech = doc.add_paragraph()
    p_tech.paragraph_format.space_before = Pt(5)
    p_tech.paragraph_format.space_after = Pt(15)
    run_tech = p_tech.add_run("Meeting strict online processing constraints: <5 mins execution and <16 GB RAM.")
    run_tech.font.name = 'Calibri'
    run_tech.font.size = Pt(14)
    run_tech.font.bold = True
    
    add_bullet(doc, "• Precomputed Sentence Embeddings:", "Candidates are embedded offline using the lightweight, high-performance 'all-MiniLM-L6-v2' model. Embeddings are stored as a 153MB dense matrix (`candidate_embeddings.npy`) on disk.")
    add_bullet(doc, "• High-Speed Vectorized Matching:", "During online execution, the Job Description is embedded in milliseconds, and semantic similarities across all 100k candidates are calculated in under 0.1 seconds using vectorized matrix multiplication.")
    add_bullet(doc, "• Single-Pass Scorer Execution:", "Candidate metadata loading and custom scorer evaluation are executed in a single-pass loop, completing the entire ranking pipeline in under 5 seconds.")
    add_bullet(doc, "• Zero API Dependencies:", "The system runs entirely offline with local SentenceTransformers weights and logic, ensuring absolute compliance with challenge guidelines.")

    # --- SLIDE 8: Recruiter Dashboard UI/UX ---
    doc.add_page_break()
    add_slide_header(doc, "7. Interactive Recruiter Dashboard")
    
    p_ui = doc.add_paragraph()
    p_ui.paragraph_format.space_before = Pt(5)
    p_ui.paragraph_format.space_after = Pt(15)
    run_ui = p_ui.add_run("A premium recruitment workspace designed to make decisions faster and clearer.")
    run_ui.font.name = 'Calibri'
    run_ui.font.size = Pt(14)
    run_ui.font.italic = True
    
    add_bullet(doc, "• Glassmorphic Dark Design:", "Premium visual style utilizing a curated, cohesive dark mode color scheme (Indigo, Emerald, Charcoal) and Outfit typography.")
    add_bullet(doc, "• Interactive Filter Workspace:", "Dynamic sidebar allowing recruiters to filter by experience (min-slider), notice period, local Noida/Pune status, and activity intervals.")
    add_bullet(doc, "• Score Breakdown & Explainability:", "Displays individual candidate sub-scores (semantic, skills, experience, engagement) and lists explicit strengths, weaknesses, and missing skills.")
    add_bullet(doc, "• Career Timeline & Skill Gaps:", "Visualizes candidates' career timelines dynamically and checks their skills against required/preferred categories with check/cross icons.")

    # --- SLIDE 9: Evaluation & Verification ---
    doc.add_page_break()
    add_slide_header(doc, "8. Verification & Validation Metrics")
    
    p_eval = doc.add_paragraph()
    p_eval.paragraph_format.space_before = Pt(5)
    p_eval.paragraph_format.space_after = Pt(15)
    run_eval = p_eval.add_run("Rigorous testing to guarantee ranking reliability, consistency, and compliance.")
    run_eval.font.name = 'Calibri'
    run_eval.font.size = Pt(14)
    run_eval.font.bold = True
    
    add_bullet(doc, "• Auto-Validator Alignment:", "Verified that candidate rankings contain exactly 100 entries in `submission.csv`, sorted in strictly non-increasing order of scores.")
    add_bullet(doc, "• Deterministic Tie-breaking:", "Resolves equal scores using candidate ID in ascending alphabetical order, ensuring matching runs produce identical shortlists every time.")
    add_bullet(doc, "• 0% Honeypot Kurz-Verification:", "Verified that honeypot candidates are filtered out of the top 100 recomendations, ensuring zero recruiter-facing quality regressions.")
    add_bullet(doc, "• Schema Verification:", "Runs the provided `validate_submission.py` to guarantee that output columns ('candidate_id', 'rank', 'score', 'reasoning') conform to requirements.")

    # --- SLIDE 10: Conclusion & Next Steps ---
    doc.add_page_break()
    add_slide_header(doc, "9. Future Roadmap & Conclusion")
    
    p_road = doc.add_paragraph()
    p_road.paragraph_format.space_before = Pt(5)
    p_road.paragraph_format.space_after = Pt(15)
    run_road = p_road.add_run("Key takeaways and long-term features for SmartHire AI:")
    run_road.font.name = 'Calibri'
    run_road.font.size = Pt(14)
    run_road.font.italic = True
    
    add_bullet(doc, "• Summary:", "SmartHire AI successfully bridges the gap between structured signals and semantic context, identifying the absolute best fits from 100k records in seconds.")
    add_bullet(doc, "• Two-Stage LLM Re-ranking:", "For future scale, we can feed the top 500 candidates into a local, quantized LLM (e.g. Llama-3-8B-Instruct) to perform advanced instruction-based re-ranking.")
    add_bullet(doc, "• Fine-tuned Embeddings:", "Contrastive fine-tuning of the embedding model using historical hire/rejection datasets can align semantic embeddings closer to actual company-specific preferences.")
    add_bullet(doc, "• Hybrid Search Indexing:", "Implementing Qdrant or Milvus in the backend to enable vector and lexical hybrid search dynamically as new candidates are added.")

    # Save DOCX
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(base_dir, "data", "presentation.docx")
    doc.save(docx_path)
    print(f"Presentation saved to {docx_path}")
    
    # Try converting to PDF using Word COM
    pdf_path = os.path.join(base_dir, "data", "presentation.pdf")
    print(f"Attempting to compile DOCX to PDF at {pdf_path} using Microsoft Word COM...")
    
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc_abs = os.path.abspath(docx_path)
        pdf_abs = os.path.abspath(pdf_path)
        
        doc_obj = word.Documents.Open(doc_abs)
        # 17 represents wdFormatPDF
        doc_obj.SaveAs(pdf_abs, FileFormat=17)
        doc_obj.Close()
        word.Quit()
        print("PDF conversion completed successfully!")
    except Exception as e:
        print(f"Python win32com conversion failed: {e}")
        print("Falling back to PowerShell for COM conversion...")
        
        cmd = (
            f'powershell -Command "$word = New-Object -ComObject Word.Application; '
            f'$word.Visible = $false; '
            f'$doc = $word.Documents.Open(\'{os.path.abspath(docx_path)}\'); '
            f'$doc.SaveAs(\'{os.path.abspath(pdf_path)}\', 17); '
            f'$doc.Close(); $word.Quit();"'
        )
        ret = os.system(cmd)
        if ret == 0:
            print("PowerShell PDF conversion completed successfully!")
        else:
            print("PowerShell PDF conversion failed.")

if __name__ == "__main__":
    main()
